import random

import time
from datetime import datetime, timedelta
import os
import sys

from tkinter import *

from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import RGBColor
from docx.shared import Inches

import xlrd
import openpyxl

from collections import Counter
import re

from bs4 import BeautifulSoup
import requests

#______________________________________________________________________________
#TEXTS

#start
b1 = "The Casino team expressed an intention to check \"xxx's\" overall gameplay. Therefore, this matter is investigated further in the report."
b1b = "Initial query from the Casino raised a concern about the customer\'s \"xxx\" overall activity. Therefore, the following investigation was performed."

b2 = "The team expressed an intention to check \"xxx's\" gameplay during the period from date. Therefore, this matter is investigated further in the report."
b2b = "Initial query from the licensee raised a concern regarding the player\'s sessions from date. Therefore, the following investigation was performed."

b3 = "The team expressed an intention to check \"xxx's\" gameplay for date session. Therefore, this matter is investigated further in the report."
b3b = "Initial request from the licensee expressed a concern about the user\'s gameplay on date. Therefore, the following investigation was performed."

#dealer preference
c1 = 'Also, during investigated period the customer held games regardless of a dealer.'
c2 = 'What is more, games were played regardless of the present dealer.'
c3 = 'Moreover, the player did not follow any dealer in particular.'

#vids
e2 = "What is more, all video records and logs of games that generated the most profit were checked, to see if any deviations from regular procedures took place, as a result, all the procedures were followed correctly."
e2b = 'The video records and logs for most profitable rounds were investigated. After necessary checks, no irregularities in the logs were observed and all of the actions performed by the dealer were done according to set procedures.'
e2c = 'Additionally, the video record invetigation of the player\'s most profitable games was performed, no irregularities were established nor any conditions that may have developed that could affect the fairness of the game were observed.'
e2d = 'Additionally, after going through the video records and logs for the players\' most profitable games, all procedures were followed correctly and no conditions that may have developed to be affecting the fairness of the game were observed.'

#opposite b
e3 = "The customer was also investigated for any signs of an opposite betting i.e. two accounts are playing in the same games wagering the same or similar amount on opposite spots so that makes one player lose and the second win within Evolution Gaming Network. The investigation revealed no correlation of \"xxx\'s\" bets with any other players bets."
e3b = "In addition, \"xxx\'s\" gameplay was investigated in order to identify whether there were other accounts within Evolution Gaming network who were participating in the same games by placing equal or similar bet amounts on opposite spots; accordingly, what makes one user lose and the second win. After necessary checks, no other accounts that could match the described betting pattern were observed."
e3c = 'In additional, \"xxx\'s\" activity was checked for signs of opposite betting, in order to identify whether there were other accounts within Evolution Gaming network who were participating in the same games by placing equal or similar bet amounts on opposite spots. After necessary checks, no other accounts that could match the described betting pattern were observed.'

#fin
f1 = "To sum up, the customer\'s gaming activity does not reveal suspicious trends or signs of fraudulent actions. After all facts observed it is considered, that the customer is not a threat to the Live Casino environment."
f2 = "In conclusion, the user\'s gameplay does not reveal suspicious trends and no signs of fraudulent activity were observed. \"xxx\" is not a threat to Live Casino environment."
f3 = 'To summarize, the customer\'s gameplay does not reveal suspicions trends and no signs of illicit activity were detected. All things considered, \"xxx\" is not a threat to Live Casino environment.'

#______________________________________________________________________________

root = Tk()
root.title('General Player Reporting')
root.geometry('400x170')

#_____________________________________________________________________________
#MANIN FUNCTION

def reporting_main(event):
    
    global b1
    global b1b
    global b2
    global b2b
    global b3
    global b3b
    global e3
    global e3b
    global e3c
    global f1
    global f2
    global f3
    
    document = Document('General_reporting.docx')
    
    #Essential functions before main function
    
    #Table cell name proclamation
    
    table0 = document.tables[0]   # info about the player
    
    cell011 = table0.cell(1,1)    # operator ID
    
    cell021 = table0.cell(2,1)    # screen name
    
    cell031 = table0.cell(3,1)    # UID
    
    cell040 = table0.cell(4,0)    # agg. win/loss
    
    cell041 = table0.cell(4,1)    # profit
    
    cell051 = table0.cell(5,1)    # turnover
    
    cell061 = table0.cell(6,1)    # Margin
    
    cell071 = table0.cell(7,1)    # date of enrollment
    
    cell081 = table0.cell(8,1)
    #
    table1 = document.tables[1]   # Body of a text
    
    cell110 = table1.cell(1,0)
    #
    table2 = document.tables[2]   # Conclusion
    
    cell210 = table2.cell(1,0)
    
    #Adding 1st/2nd Names + date on the cover page + Timeframe of analysis
    
    run0 = document.paragraphs[13].add_run(os.environ['USERNAME'])
    font0 = run0.font
    font0.size = Pt(15)
    
    if os.environ['USERNAME'] == 'Aleksandrs':
        run01 = document.paragraphs[13].add_run(' Karsonis')
        font01 = run01.font
        font01.size = Pt(14)    
        if os.environ['USERNAME'] == 'Arturs':
            run02 = document.paragraphs[13].add_run(' Lusis')
            font02 = run01.font
            font02.size = Pt(14)
            if os.environ['USERNAME'] == 'Alina':
                run02 = document.paragraphs[13].add_run(' Heifeca')
                font02 = run01.font
                font02.size = Pt(14)        
    
    run2 = document.paragraphs[15].add_run(time.strftime("%d.%m.%y"))
    font2 = run2.font
    font2.size = Pt(14)
    
    dates = dates_entry.get()
    if dates == '':
        run3 = document.paragraphs[14].add_run('All history')
        font2 = run3.font
        font2.size = Pt(14)
    else:
        run3 = document.paragraphs[14].add_run(
            '\nAll history; in-depth for ' + dates)
        font2 = run3.font
        font2.size = Pt(14)    
    
    #__________________________________________________________________________
    #Logging in
    try:
        with requests.Session() as c:
            url = 'https://sbo.ezugi.com/office.php?page=login'
            USERNAME = 'akarsonis'
            PASSWORD = 'ezugi123456'
            c.get(url)
            login_data = dict(username=USERNAME, password=PASSWORD, 
                              language_view='english', submit='Login')
            c.post(url, 
                   data=login_data, 
                   headers={'Referer': 
                            'https://sbo.ezugi.com/office.php?page=login'})
            page = c.get(link_entry.get())
        
        soup = BeautifulSoup(page.content, 'html.parser')
        
        #__________________________________________________________________________
        #Mining Player info
        #Parsing operator ID
    
        operator_id = str(soup.find('img', {'height' : '20'}))
        try:
            operator_id = ((operator_id.split('ID: '))[1].split(' " width="20"/>')[0])
        except IndexError:
            print('incorrect link provided')
        operator_id = operator_id.replace('\n', '')
        operator_id = operator_id.replace('  ', ' ')
        operator_id = operator_id[:-1]
        
        #Parsing uid
        
        uid = str(soup.find('input', {'name' : 'PlayerDisplay'}))
        uid = uid.split('" type="text" value="')[1].split('"/>')[0]
        
        #Parsing screen name2
        
        screen_name = str(soup.find('img', {'width' : '25'}))
        screen_name = screen_name.split(' - ')[1].split(' Operator: ')[0]
        screen_name = screen_name.replace('\n', '')
        
        #Parsing currency
        
        currency = str(soup.find('td', {'class' : 'grid_cell SessionCurrency'}))
        currency = currency.split('title="')[1].split(
            '"><div id="SessionCurrency"')[0]
        
        #Parsing turnover
        
        turnover = str(soup.find('td', {'class' : 'grid_cell BetSum'}))
        turnover = turnover.split('title="')[1].split('"><div id="BetSum"')[0]
        
        #Parsing payoff
        
        payoff = str(soup.find('td', {'class' : 'grid_cell WinSum'}))
        payoff = payoff.split('title="')[1].split('"><div id="WinSum"')[0]
        
        #Calculating net result + margin
        
        net = float(payoff) - float(turnover)
        margin = float(net) / float(turnover) * 100
        
        #Formatting values
        
        turnover = "{:5,.2f}".format(float(turnover))
        net_formated = "{:5,.2f}".format(net)
        margin = "{:5,.2f}".format(margin)
        
        #__________________________________________________________________________
        #Writting into cells
        
        cell011.text = operator_id
        cell021.text = screen_name
        cell031.text = uid
        cell061.text = margin + '%'
        
        if net > 0:
            cell040.paragraphs[0].add_run('Aggregated win:').bold = True
        elif net == 0:
            cell040.paragraphs[0].add_run('Result:').bold = True
        else:
            cell040.paragraphs[0].add_run('Aggregated loss:').bold = True
        
        #If main currency not EURO
        
        if currency != 'EUR':
            turnover_eur = str(soup.find('td', {'class' : 'grid_cell BetUSD'}))
            turnover_eur = turnover_eur.split('title="')[1].split(
                '"><div id="BetUSD"')[0]
            
            payoff_eur = str(soup.find('td', {'class' : 'grid_cell WinUSD'}))
            payoff_eur = payoff_eur.split('title="')[1].split(
                '"><div id="WinUSD"')[0]
            
            net_eur = float(payoff_eur) - float(turnover_eur)
            
            net_eur = "{:5,.2f}".format(net_eur)
            turnover_eur = "{:5,.2f}".format(float(turnover_eur))
            
            cell041.text = (currency + ' ' + str(net_formated)) + ' | ' + (
                'EUR ' +  str(net_eur))
            cell051.text = (currency + ' ' + str(turnover)) + ' | ' + (
                'EUR ' + str(turnover_eur))
        else:
            cell041.text = currency + ' ' + str(net_formated)
            cell051.text = currency + ' ' + str(turnover)        
        
        #First login
        
        first_login = str(soup.find(
            'td', {'class' : 'grid_cell PlayerFirstLogin'}))
        try:
            first_login = ((first_login.split('hidden;">'))[1].split(' ')[0])
        except IndexError:
            print('incorrect link provided')
        
        year_login = '.' + first_login[0:4] 
        month_login = '.' + first_login[5:7]
        day_login = first_login[8:]
        
        propper_first_login = day_login + month_login + year_login
        cell071.text = propper_first_login
        
        #__________________________________________________________________________
        #ADDING PARAGRAPHS
        #First paragraph
            
        b1 = b1.replace('date', dates)
        b2 = b2.replace('date', dates)
        b2b = b2b.replace('date', dates)
        b3 = b3.replace('date', dates)
        b3b = b3b.replace('date', dates)
        
        b1 = b1.replace('xxx', screen_name)
        b1b = b1b.replace('xxx', screen_name)
        b2 = b2.replace('xxx', screen_name)
        b2b = b2b.replace('xxx', screen_name)
        b3 = b3.replace('xxx', screen_name)
        b3b = b3b.replace('xxx', screen_name)
        
        if len(dates) > 23:
            cell110.paragraphs[0].add_run(random.choice([b2, b2b]))
        elif len(dates) < 24 and len(dates) > 1:
            cell110.paragraphs[0].add_run(random.choice([b3, b3b]))
        elif len(dates) < 1:
            cell110.paragraphs[0].add_run(random.choice([b1, b1b]))
            
        cell110.add_paragraph('\n')
            
        #Dealer preference paragraph
        
        if depref_var.get() == 'Yes':
            cell110.add_paragraph(
                random.choice([c1, c2, c3]), style = 'No Spacing')
        else:
            pass    
        
        #Video records paragraph
        
        if video_var.get() == 'Yes':
            cell110.add_paragraph(
                random.choice([e2, e2b, e2c, e2d]), style = 'No Spacing')
        else:
            pass
        
        #Opposite betting paragraph
        
        e3 = e3.replace('xxx', screen_name)
        e3b = e3b.replace('xxx', screen_name)
        e3c = e3c.replace('xxx', screen_name)    
        
        if opposite_var.get() == 'Yes':
            cell110.add_paragraph(
                random.choice([e3, e3b, e3c]), style = 'No Spacing')
        else:
            pass
        
        #Final paragraph
        
        f2 = f2.replace('xxx', screen_name)
        f3 = f3.replace('xxx', screen_name)
        cell210.paragraphs[0].add_run(random.choice([f1, f2, f3]))
    except IndexError:
        generate_text_var.set("WRONG LINK! TRY AGAIN!")
    else:
        generate_text_var.set("LINK IS GOOD!")
    
    #__________________________________________________________________________
    #Save
    
    try:
        document.save(
            'Risk Assessment Report ' + str(screen_name) + '.docx')
    except PermissionError:
        document.save(
            'Risk Assessment Report ' + str(screen_name) + '(1).docx')    

#_____________________________________________________________________________
#TKINTER ARCHITECTURE
#Date

dates_label = Label(
    root, text = 'Timeframe of analysis (ALL HISTORY by default)')
dates_label.grid(row=0, column=0, sticky=W)

dates_entry = Entry(root)
dates_entry.grid(row=0, column=1, sticky=E)

#SBO link

link_label = Label(root, text = 'SBO link')
link_label.grid(row=1, column=0, sticky=W)

link_entry = Entry(root)
link_entry.grid(row=1, column=1, sticky=E)

#Dealer preference

depref_label = Label(root, text = 'Dealer preference checked?')
depref_label.grid(row=2, column=0, sticky=W)

depref_list = ['Yes', 'No']
depref_var = StringVar(root)
depref_var.set('Yes')

depref_optionmenu = OptionMenu(root, depref_var, *depref_list)
depref_optionmenu.grid(row=2, column=1, sticky=E)

#Video record check

video_label = Label(root, text = 'Video records checked?')
video_label.grid(row=3, column=0, sticky=W)

video_list = ['Yes', 'No']
video_var = StringVar(root)
video_var.set('Yes')

video_optionmenu = OptionMenu(root, video_var, *video_list)
video_optionmenu.grid(row=3, column=1, sticky=E)

#Oppotise betting check

opposite_label = Label(root, text = 'Opposite betting checked?')
opposite_label.grid(row=4, column=0, sticky=W)

opposite_list = ['Yes', 'No']
opposite_var = StringVar(root)
opposite_var.set('Click to choose')

opposite_optionmenu = OptionMenu(root, opposite_var, *opposite_list)
opposite_optionmenu.grid(row=4, column=1, sticky=E)

#Generate

generate_text_var = StringVar(root)
generate = Button(root, 
                  textvariable=generate_text_var)
generate_text_var.set("Generate")

generate.grid(row=5, column=0, columnspan=2)
generate.bind('<Button-1>', reporting_main)

root.mainloop()